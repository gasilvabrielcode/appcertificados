from docx import Document  # manipulador de arquivos
import os  # sistema operacional, navegar em pastas e poder excluir/editar arquivos
import mysql.connector  # faz conexão com banco mysql
from pdf_mail import sendpdf  # faz o envio de pdf por e-mail
from colorama import Fore  # modulo para adicionar cores
from datetime import datetime  # modulo para puxar datas
import subprocess # converter pdf

#  endereço sql para fazer conexão com banco
mydb = mysql.connector.connect(
    host="177.47.183.75",
    user="tech_consulta",
    password="~pt780w3L",
    database="icetran_oraculo"
)

#  SELECT para puxar as variáveis
myCursor = mydb.cursor()
idusuario = '100'
myCursor.execute("SELECT nome, senha, idcidade FROM usuarios_adm WHERE idusuario = {}".format(idusuario))
myresults = myCursor.fetchall()

#  colocar as variáveis do banco no código
for myresult in myresults:
    x = datetime.now()
    idcurso = input('ID CURSO: ')
    nome = input('NOME: ').upper().strip()
    email = input('Email de {}: '.format(nome))
    cpf = (input("CPF: "))


    def formatcpf(cpf):
        vezes = 0
        novo = ""
        for quantidade in range(11):
            numero = cpf[quantidade]
            novo += numero
            vezes += 1
            if quantidade == 8:
                novo += "-"
                vezes -= 3
            if vezes == 3:
                novo += "."
                vezes -= 3
        return novo


    cnh = input('CNH: ')
    renach = input('RENACH: ')
    categoria = input('CATEGORIA: ')
    certificado = input('MATRÍCULA: ')
    match x.month:
        case 1:
            mes = 'Janeiro'
        case 2:
            mes = 'Fevereiro'
        case 3:
            mes = 'Março'
        case 4:
            mes = 'Abril'
        case 5:
            mes = 'Maio'
        case 6:
            mes = 'Junho'
        case 7:
            mes = 'Julho'
        case 8:
            mes = 'Agosto'
        case 9:
            mes = 'Setembro'
        case 10:
            mes = 'Outubro'
        case 11:
            mes = 'Novembro'
        case 12:
            mes = 'Dezembro'
        case _:
            mes = ''
    data = ('{} de {} de {}'.format(x.day, mes, x.year))
    inicio = input('INICIO: ').replace(".", "/")
    fim = input('FIM: ').replace(".", "/")
    validade = fim.split("/")
    val_dia = int(validade[0])
    val_mes = int(validade[1])
    val_ano = int(validade[2])
    validade = datetime((val_ano + 5), val_mes, val_dia)
    validade = datetime.strftime(validade, "%d/%m/%Y")

    nota_total = input('NOTA: ')
    nota1 = float(nota_total)
    nota = str(round((nota1 * 10) / 30, 1))
    aproveitamento_total = (nota1 * 100) / 30
    aproveitamento = str(round(aproveitamento_total)) + '%'
    registro = input('REGISTRO: ')
    folha = input('FOLHA: ')
    livro = input('LIVRO: ')
    validacao = input('CÓDIGO CERTIFICADO: ')
    direcao = input('DIREÇÃO DEFENSIVA: ')
    primeiros = input('PRIMEIROS SOCORROS: ')
    mecanica = input('MECÂNICA BÁSICA: ')
    relacoes = input('RELAÇÕES HUMANAS: ')
    conceitos = input('PRINCIPAIS CONCEITOS: ')
    legislacao = input('LEGISLAÇÃO: ')

#  puxar os templates de acordo com o id do curso
if idcurso == '1':
    certificado_template = 'templates/taxista_pe.docx'
if idcurso == '2':
    certificado_template = 'templates/taxista_goytacazes.docx'
if idcurso == '3':
    certificado_template = 'templates/rci_rs.docx'
if idcurso == '4':
    certificado_template = 'templates/mopp_sc.docx'
if idcurso == '173':
    certificado_template = 'templates/emergencia_sc.docx'


#  função para gerar o arquivo resultado
def main():
    template_file_path = certificado_template
    output_file_path = 'result.docx'

    #  nome das variáveis para colocar dentro do pdf e a função
    variables = {
        "${NOME}": nome,
        "${CPF}": formatcpf(cpf),
        "${CNH}": cnh,
        "${RENACH}": renach,
        "${CATEGORIA}": categoria,
        "${ID}": certificado,
        "${CERTIFICADO}": certificado,
        "${INICIO}": inicio,
        "${FIM}": fim,
        "${VALIDADE}": validade,
        "${NOTA}": nota,
        "${DATA}": data,
        "${APROVEITAMENTO}": aproveitamento,
        "${REGISTRO}": registro,
        "${FOLHA}": folha,
        "${LIVRO}": livro,
        "${VALIDACAO}": validacao,
        "${DIRECAO}": direcao,
        "${PRIMEIROS}": primeiros,
        "${MECANICA}": mecanica,
        "${RELACOES}": relacoes,
        "${CONCEITOS}": conceitos,
        "${LEGISLACAO}": legislacao,
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()

#  conversor de docx para pdf e também para jogar ao diretório /certificados
output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'result.docx'])
os.rename('/home/gabriel/projects/apps_py/appcertificados/result.pdf',
          '/home/gabriel/projects/apps_py/appcertificados/certificados/{}.pdf'.format(nome))

#  input e função para envio de e-mail
print('[1] - ', Fore.GREEN + 'SIM')
print(Fore.RESET + '[2] - ', Fore.RED + 'NÃO')
reposta = str(input(Fore.RESET + 'Deseja enviar o Certificado de {} por E-mail?\n'.format(nome)))
if reposta == '1':
    email_remetente = 'gabriel.silva@ibrep.com.br'
    email_destinatario = email
    email_senha = ''
    email_assunto = 'SEGUNDA VIA CERTIFICADO ICETRAN'
    email_corpo = 'Olá {}, segue em anexo a sua segunda via certificado!\n' \
                  'Nós da equipe ICETRAN, desejamos um bom dia!'.format(nome)
    arquivo = '{}'.format(nome)
    arquivo_local = "/home/gabriel/projects/apps_py/appcertificados/certificados"
    envio = sendpdf(email_remetente,
                    email_destinatario,
                    email_senha,
                    email_assunto,
                    email_corpo,
                    arquivo,
                    arquivo_local)
    envio.email_send()
    print(Fore.LIGHTGREEN_EX + 'ENVIADO COM SUCESSO!')
else:
    print('Obrigado!')

os.remove('result.docx')
