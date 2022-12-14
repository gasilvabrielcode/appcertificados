#!/usr/bin/python
import tkinter  # biblioteca de ui
import customtkinter  # biblioteca de ui customizada
from docx import Document  # manipulador de arquivos
import os  # sistema operacional, navegar em pastas e poder excluir/editar arquivos
import mysql.connector  # faz conexão com banco mysql
from pdf_mail import sendpdf  # faz o envio de pdf por e-mail
from datetime import datetime  # modulo para puxar datas
import subprocess  # converter pdf

customtkinter.set_appearance_mode("dark")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"
switch_value = True

x = datetime.now()
app = customtkinter.CTk()
app.geometry("470x450")
app.title("Tech App")
app2 = customtkinter.CTk()
app2.geometry("250x720")
app2.title("Temp")


#  endereço sql para fazer conexão com banco
# mydb = mysql.connector.connect(
#     host="177.47.183.75",
#     user="tech_consulta",
#     password="~pt780w3L",
#     database="icetran_oraculo"
# )

#  SELECT para puxar as variáveis
# myCursor = mydb.cursor()
# idusuario = '100'
# myCursor.execute("SELECT nome, senha, idcidade FROM usuarios_adm WHERE idusuario = {}".format(idusuario))
# myresults = myCursor.fetchall()


def button_callback():
    def formatcpf(cpf):
        vezes = 0
        novo = ""
        for quantidade in range(11):
            numero = cpf[quantidade]
            novo += numero
            vezes += 1
            if quantidade == 8:
                novo += "-"
                vezes -= 3
            if vezes == 3:
                novo += "."
                vezes -= 3
        return novo

    NOME = str(entry_NOME.get()).upper()
    CPF = entry_CPF.get()
    CPF = str(formatcpf(CPF))
    CNH = str(entry_CNH.get())
    RENACH = str(entry_RENACH.get())
    CATEGORIA = str(entry_CATEGORIA.get())
    CERTIFICADO = str(entry_ID.get())
    INICIO = str(entry_INICIO.get())
    FIM = str(entry_FIM.get())
    REGISTRO = str(entry_REGISTRO.get())
    FOLHA = str(entry_FOLHA.get())
    LIVRO = str(entry_LIVRO.get())
    VALIDACAO = str(entry_VALIDACAO.get())
    DIRECAO = str(entry_DIRECAO.get())
    PRIMEIROS = str(entry_PRIMEIROS.get())
    MECANICA = str(entry_MECANICA.get())
    RELACOES = str(entry_RELACOES.get())
    CONCEITOS = str(entry_CONCEITOS.get())
    LEGISLACAO = str(entry_LEGISLACAO.get())

    match x.month:
        case 1:
            mes = 'Janeiro'
        case 2:
            mes = 'Fevereiro'
        case 3:
            mes = 'Março'
        case 4:
            mes = 'Abril'
        case 5:
            mes = 'Maio'
        case 6:
            mes = 'Junho'
        case 7:
            mes = 'Julho'
        case 8:
            mes = 'Agosto'
        case 9:
            mes = 'Setembro'
        case 10:
            mes = 'Outubro'
        case 11:
            mes = 'Novembro'
        case 12:
            mes = 'Dezembro'
        case _:
            mes = ''
    data = ('{} de {} de {}'.format(x.day, mes, x.year))
    validade = entry_FIM.get().split(".")
    val_dia = int(validade[0])
    val_mes = int(validade[1])
    val_ano = int(validade[2])
    validade = datetime((val_ano + 5), val_mes, val_dia)
    validade = datetime.strftime(validade, "%d/%m/%Y")

    nota_total = entry_NOTA.get()
    nota1 = float(nota_total)
    nota = str(round((nota1 * 10) / 30, 1))
    aproveitamento_total = (nota1 * 100) / 30
    aproveitamento = str(round(aproveitamento_total)) + '%'

    progressbar_1.set(0.25)

    variables = {
        "${NOME}": NOME.upper(),
        "${CPF}": CPF,
        "${CNH}": CNH,
        "${RENACH}": RENACH,
        "${CATEGORIA}": CATEGORIA,
        "${ID}": CERTIFICADO,
        "${CERTIFICADO}": CERTIFICADO,
        "${INICIO}": INICIO.replace(".", "/"),
        "${FIM}": FIM.replace(".", "/"),
        "${VALIDADE}": validade,
        "${NOTA}": nota,
        "${DATA}": data,
        "${APROVEITAMENTO}": aproveitamento,
        "${REGISTRO}": REGISTRO,
        "${FOLHA}": FOLHA,
        "${LIVRO}": LIVRO,
        "${VALIDACAO}": VALIDACAO,
        "${DIRECAO}": DIRECAO,
        "${PRIMEIROS}": PRIMEIROS,
        "${MECANICA}": MECANICA,
        "${RELACOES}": RELACOES,
        "${CONCEITOS}": CONCEITOS,
        "${LEGISLACAO}": LEGISLACAO,
    }
    certificado_template = ''
    if optionmenu_Curso.get() == "Taxista" and optionmenu_Estado.get() == "Pernambuco":
        certificado_template = 'templates/taxista_pe.docx'
    if optionmenu_Curso.get() == "Taxista" and optionmenu_Estado.get() == "RJ - Goytacazes":
        certificado_template = 'templates/taxista_goytacazes.docx'
    if optionmenu_Curso.get() == "Reciclagem" and optionmenu_Estado.get() == "Santa Catarina":
        certificado_template = 'templates/rci_rs.docx'
    if optionmenu_Curso.get() == "MOPP" and optionmenu_Estado.get() == "Santa Catarina":
        certificado_template = 'templates/mopp_sc.docx'
    if optionmenu_Curso.get() == "Emergência" and optionmenu_Estado.get() == "Santa Catarina":
        certificado_template = 'templates/emergencia_sc.docx'

    def main():
        template_file_path = certificado_template
        output_file_path = 'result.docx'

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)

    def replace_text_in_paragraph(paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)

    progressbar_1.set(0.50)

    main()

    #  conversor de docx para pdf e também para jogar ao diretório /certificados
    output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'result.docx'])
    os.rename('./result.pdf',
              './certificados/{}.pdf'.format(NOME))

    progressbar_1.set(0.75)
    entry_NOME.delete(0, 10000000)
    entry_CPF.delete(0, 10000000)
    entry_CNH.delete(0, 10000000)
    entry_ID.delete(0, 10000000)
    entry_EMAIL.delete(0, 10000000)
    entry_INICIO.delete(0, 10000000)
    entry_FIM.delete(0, 10000000)
    entry_RENACH.delete(0, 10000000)
    entry_CATEGORIA.delete(0, 10000000)
    entry_VALIDACAO.delete(0, 10000000)
    entry_REGISTRO.delete(0, 10000000)
    entry_NOTA.delete(0, 10000000)
    entry_FOLHA.delete(0, 10000000)
    entry_LIVRO.delete(0, 10000000)
    entry_DIRECAO.delete(0, 10000000)
    entry_PRIMEIROS.delete(0, 10000000)
    entry_CONCEITOS.delete(0, 10000000)
    entry_RELACOES.delete(0, 10000000)
    entry_MECANICA.delete(0, 10000000)
    entry_LEGISLACAO.delete(0, 10000000)

    if checkbox_1.get() == 1:
        email_remetente = ''  # insira email
        email_destinatario = entry_EMAIL.get()
        email_senha = ''  # insira senha do email
        email_assunto = 'SEGUNDA VIA CERTIFICADO ICETRAN'
        email_corpo = 'Olá {}, segue em anexo a sua segunda via certificado!\n' \
                      'Nós da equipe ICETRAN, desejamos um bom dia!'.format(NOME)
        arquivo = '{}'.format(NOME)
        arquivo_local = "/home/gabriel/projects/apps_py/appcertificados/certificados"
        envio = sendpdf(email_remetente,
                        email_destinatario,
                        email_senha,
                        email_assunto,
                        email_corpo,
                        arquivo,
                        arquivo_local)
        envio.email_send()
        progressbar_1.set(1)
        enviado = customtkinter.CTkLabel(master=frame_1, justify=tkinter.LEFT, text='E-mail Enviado com Sucesso!')
        enviado.pack(pady=10, padx=10)

    else:
        obrigado = customtkinter.CTkLabel(master=frame_1, justify=tkinter.LEFT, text='Obrigado!')
        obrigado.pack(pady=10, padx=10)
        progressbar_1.set(1)
    os.remove('result.docx')


# def button_callback2():
#
#     global switch_value
#     if switch_value:
#         customtkinter.set_appearance_mode("Light")
#         switch_value = False
#     else:
#         customtkinter.set_appearance_mode("dark")
#         switch_value = True


frame_1 = customtkinter.CTkFrame(master=app)
frame_1.pack(pady=20, padx=60, fill="both", expand=True)

label_1 = customtkinter.CTkLabel(master=frame_1, justify=tkinter.LEFT, text='Gerar Segunda-Via')
label_1.pack(pady=10, padx=10)

entry_CPF = customtkinter.CTkEntry(master=frame_1, placeholder_text="CPF")
entry_CPF.pack(pady=10, padx=10)

optionmenu_Estado = customtkinter.CTkOptionMenu(frame_1, values=["Pernambuco", "Santa Catarina",
                                                                 "Rio de Janeiro (Goytacazes)"])
optionmenu_Estado.pack(pady=10, padx=10)
optionmenu_Estado.set("Selecione o Estado")

optionmenu_Curso = customtkinter.CTkOptionMenu(frame_1, values=["Taxista", "Reciclagem", "MOPP", "Emergência"])
optionmenu_Curso.pack(pady=10, padx=10)
optionmenu_Curso.set("Selecione o Curso")

checkbox_1 = customtkinter.CTkCheckBox(master=frame_1, text="Enviar E-mail")
checkbox_1.pack(pady=10, padx=10)

entry_EMAIL = customtkinter.CTkEntry(master=frame_1, placeholder_text="E-mail")
entry_EMAIL.pack(pady=10, padx=10)

button_1 = customtkinter.CTkButton(master=frame_1, command=button_callback, text='Gerar Certificado')
button_1.pack(pady=10, padx=10)

# button_2 = customtkinter.CTkButton(master=frame_1, command=button_callback2, text='Modo', width=10)
# button_2.pack(pady=10, padx=10)
# button_2.place(relx=0.03)

progressbar_1 = customtkinter.CTkProgressBar(master=frame_1)
progressbar_1.pack(pady=10, padx=10)
progressbar_1.set(0)

frame_2 = customtkinter.CTkFrame(master=app2)
frame_2.pack(fill="both", expand=True)

entry_NOME = customtkinter.CTkEntry(master=frame_2, placeholder_text="Nome")
entry_NOME.pack(pady=10, padx=10)
entry_NOME.place(anchor=tkinter.W, rely=0.1, relx=0.05)

entry_CNH = customtkinter.CTkEntry(master=frame_2, placeholder_text="CNH")
entry_CNH.pack(pady=10, padx=10)
entry_CNH.place(anchor=tkinter.W, rely=0.15, relx=0.05)

entry_RENACH = customtkinter.CTkEntry(master=frame_2, placeholder_text="RENACH")
entry_RENACH.pack(pady=10, padx=10)
entry_RENACH.place(anchor=tkinter.W, rely=0.2, relx=0.05)

entry_CATEGORIA = customtkinter.CTkEntry(master=frame_2, placeholder_text="Categoria")
entry_CATEGORIA.pack(pady=10, padx=10)
entry_CATEGORIA.place(anchor=tkinter.W, rely=0.25, relx=0.05)

entry_ID = customtkinter.CTkEntry(master=frame_2, placeholder_text="Matrícula")
entry_ID.pack(pady=10, padx=10)
entry_ID.place(anchor=tkinter.W, rely=0.3, relx=0.05)

entry_INICIO = customtkinter.CTkEntry(master=frame_2, placeholder_text="Data de Início")
entry_INICIO.pack(pady=10, padx=10)
entry_INICIO.place(anchor=tkinter.W, rely=0.35, relx=0.05)

entry_FIM = customtkinter.CTkEntry(master=frame_2, placeholder_text="Data de Fim")
entry_FIM.pack(pady=10, padx=10)
entry_FIM.place(anchor=tkinter.W, rely=0.4, relx=0.05)

entry_NOTA = customtkinter.CTkEntry(master=frame_2, placeholder_text="Nota")
entry_NOTA.pack(pady=10, padx=10)
entry_NOTA.place(anchor=tkinter.W, rely=0.45, relx=0.05)

entry_REGISTRO = customtkinter.CTkEntry(master=frame_2, placeholder_text="Registro")
entry_REGISTRO.pack(pady=10, padx=10)
entry_REGISTRO.place(anchor=tkinter.W, rely=0.5, relx=0.05)

entry_VALIDACAO = customtkinter.CTkEntry(master=frame_2, placeholder_text="Código Certificado")
entry_VALIDACAO.pack(pady=10, padx=10)
entry_VALIDACAO.place(anchor=tkinter.W, rely=0.55, relx=0.05)

entry_FOLHA = customtkinter.CTkEntry(master=frame_2, placeholder_text="Folha")
entry_FOLHA.pack(pady=10, padx=10)
entry_FOLHA.place(anchor=tkinter.W, rely=0.60, relx=0.05)

entry_LIVRO = customtkinter.CTkEntry(master=frame_2, placeholder_text="Livro")
entry_LIVRO.pack(pady=10, padx=10)
entry_LIVRO.place(anchor=tkinter.W, rely=0.65, relx=0.05)

entry_DIRECAO = customtkinter.CTkEntry(master=frame_2, placeholder_text="Direção Defensiva")
entry_DIRECAO.pack(pady=10, padx=10)
entry_DIRECAO.place(anchor=tkinter.W, rely=0.70, relx=0.05)

entry_PRIMEIROS = customtkinter.CTkEntry(master=frame_2, placeholder_text="Primeiros Socorros")
entry_PRIMEIROS.pack(pady=10, padx=10)
entry_PRIMEIROS.place(anchor=tkinter.W, rely=0.75, relx=0.05)

entry_MECANICA = customtkinter.CTkEntry(master=frame_2, placeholder_text="Mecânica Básica")
entry_MECANICA.pack(pady=10, padx=10)
entry_MECANICA.place(anchor=tkinter.W, rely=0.80, relx=0.05)

entry_RELACOES = customtkinter.CTkEntry(master=frame_2, placeholder_text="Relações Humanas")
entry_RELACOES.pack(pady=10, padx=10)
entry_RELACOES.place(anchor=tkinter.W, rely=0.85, relx=0.05)

entry_CONCEITOS = customtkinter.CTkEntry(master=frame_2, placeholder_text="Principais Conceitos")
entry_CONCEITOS.pack(pady=10, padx=10)
entry_CONCEITOS.place(anchor=tkinter.W, rely=0.90, relx=0.05)

entry_LEGISLACAO = customtkinter.CTkEntry(master=frame_2, placeholder_text="Legislação")
entry_LEGISLACAO.pack(pady=10, padx=10)
entry_LEGISLACAO.place(anchor=tkinter.W, rely=0.95, relx=0.05)

app.mainloop()
